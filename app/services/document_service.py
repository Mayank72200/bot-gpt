from io import BytesIO
from pathlib import Path

import numpy as np
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import Settings
from app.core.exceptions import BadRequestError, NotFoundError
from app.models.document import Document, DocumentChunk
from app.models.user import User
from app.services.embedding_service import EmbeddingService
from app.vector_store.faiss_index import FaissIndexManager


class DocumentService:
    def __init__(self, settings: Settings, embedding_service: EmbeddingService, faiss_manager: FaissIndexManager):
        self.settings = settings
        self.embedding_service = embedding_service
        self.faiss_manager = faiss_manager

    def _chunk_text(self, text: str) -> list[str]:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.settings.chunk_size_chars,
            chunk_overlap=self.settings.chunk_overlap_chars,
            length_function=len,
        )
        return splitter.split_text(text)

    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        reader = PdfReader(BytesIO(file_bytes))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    def _extract_docx_text(self, file_bytes: bytes) -> str:
        document = DocxDocument(BytesIO(file_bytes))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    def _extract_plain_text(self, file_bytes: bytes) -> str:
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1")

    def extract_text_from_uploaded_file(self, filename: str, file_bytes: bytes) -> str:
        suffix = Path(filename).suffix.lower()

        if suffix == ".pdf":
            text = self._extract_pdf_text(file_bytes)
        elif suffix == ".docx":
            text = self._extract_docx_text(file_bytes)
        elif suffix in {".txt", ".md", ".csv", ".json"}:
            text = self._extract_plain_text(file_bytes)
        else:
            raise BadRequestError(
                code="UNSUPPORTED_DOCUMENT_TYPE",
                message="Unsupported file type. Supported types: .pdf, .docx, .txt, .md, .csv, .json",
            )

        cleaned_text = text.strip()
        if not cleaned_text:
            raise BadRequestError(code="EMPTY_DOCUMENT", message="Uploaded document has no readable text.")
        return cleaned_text

    async def upload_document(self, db: Session, user_id: str, title: str, text: str) -> tuple[Document, int]:
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            raise NotFoundError("User not found.")

        if not text.strip():
            raise BadRequestError(code="INVALID_DOCUMENT", message="Document text cannot be empty.")

        document = Document(user_id=user_id, title=title)
        db.add(document)
        db.flush()

        chunks = self._chunk_text(text)
        chunk_rows: list[DocumentChunk] = []
        for idx, _ in enumerate(chunks):
            row = DocumentChunk(document_id=document.id, chunk_index=idx)
            db.add(row)
            chunk_rows.append(row)

        db.flush()

        vectors: list[np.ndarray] = []
        metadatas: list[dict] = []
        ids: list[str] = []
        for chunk_row, chunk_text in zip(chunk_rows, chunks, strict=False):
            vector = await self.embedding_service.embed_text(chunk_text)
            vectors.append(vector)
            metadatas.append(
                {
                    "user_id": user_id,
                    "document_id": document.id,
                    "chunk_id": chunk_row.id,
                    "chunk_index": chunk_row.chunk_index,
                }
            )
            ids.append(str(chunk_row.id))

        if vectors:
            self.faiss_manager.add_text_embeddings(texts=chunks, embeddings=vectors, metadatas=metadatas, ids=ids)

        try:
            db.commit()
        except Exception as exc:
            db.rollback()
            raise BadRequestError(code="DB_COMMIT_FAILED", message="Failed to save document.") from exc

        db.refresh(document)
        return document, len(chunk_rows)
